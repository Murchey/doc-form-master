import json
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


class DocxParser:
    def __init__(self, docx_path):
        self.docx_path = Path(docx_path)
        self.document = Document(docx_path)

        self.ast = {
            "metadata": {},
            "paragraphs": [],
            "tables": [],
            "images": [],
            "formulas": [],
            "styles": [],
            "sections": []
        }

    def validate_docx(self):
        if not self.docx_path.exists():
            raise FileNotFoundError(
                f"File not found: {self.docx_path}"
            )

        if self.docx_path.suffix.lower() != ".docx":
            raise ValueError(
                "Only .docx files are supported"
            )

    def parse_metadata(self):
        core = self.document.core_properties

        self.ast["metadata"] = {
            "author": core.author,
            "title": core.title,
            "created": str(core.created),
            "modified": str(core.modified)
        }

    def parse_styles(self):
        for style in self.document.styles:
            self.ast["styles"].append({
                "name": style.name,
                "type": str(style.type)
            })

    def parse_paragraphs(self):
        image_rels = self._build_image_relationship_map()

        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

        with zipfile.ZipFile(self.docx_path, "r") as zip_ref:
            media_cache = {}

            for idx, para in enumerate(self.document.paragraphs):
                paragraph_data = {
                    "id": idx,
                    "type": "paragraph",
                    "text": para.text,
                    "style": para.style.name if para.style else None,
                    "alignment": str(para.alignment),
                    "runs": []
                }

                for child in para._element:
                    tag_local = child.tag.split('}')[1] if '}' in child.tag else child.tag
                    ns_uri = child.tag.split('}')[0].strip('{') if '}' in child.tag else ''

                    if ns_uri == M_NS and tag_local in ('oMath', 'oMathPara'):
                        xml_str = etree.tostring(child).decode('utf-8')
                        math_texts = []
                        for mt in child.iter(f'{{{M_NS}}}t'):
                            if mt.text:
                                math_texts.append(mt.text)
                        paragraph_data["runs"].append({
                            "type": "formula",
                            "formula_type": tag_local,
                            "xml": xml_str,
                            "text": ''.join(math_texts)
                        })
                        continue

                    if child.tag != f'{{{W_NS}}}r':
                        continue

                    run_text = ''
                    run_bold = None
                    run_italic = None
                    run_underline = None
                    run_font_name = None
                    run_font_size = None

                    has_drawing = False

                    for sub_child in child:
                        tag_local = sub_child.tag.split('}')[1] if '}' in sub_child.tag else sub_child.tag

                        if tag_local == 't' and sub_child.text:
                            run_text += sub_child.text
                        elif tag_local == 'rPr':
                            for rpr_child in sub_child:
                                rpr_tag = rpr_child.tag.split('}')[1] if '}' in rpr_child.tag else rpr_child.tag
                                if rpr_tag == 'b':
                                    run_bold = True
                                elif rpr_tag == 'i':
                                    run_italic = True
                                elif rpr_tag == 'u':
                                    run_underline = True
                                elif rpr_tag == 'rFonts':
                                    run_font_name = rpr_child.get(f'{{{W_NS}}}eastAsia') or rpr_child.get(f'{{{W_NS}}}ascii')
                                elif rpr_tag == 'sz':
                                    sz_val = rpr_child.get(f'{{{W_NS}}}val')
                                    if sz_val:
                                        run_font_size = str(int(sz_val) // 2) + 'pt'
                        elif tag_local in ('drawing', 'object'):
                            has_drawing = True
                        elif tag_local == 'AlternateContent':
                            for ac_child in sub_child:
                                for drawing_elem in ac_child.iter():
                                    d_tag = drawing_elem.tag.split('}')[1] if '}' in drawing_elem.tag else drawing_elem.tag
                                    if d_tag == 'drawing':
                                        has_drawing = True

                    if has_drawing:
                        blip = child.findall(f'.//{{{A_NS}}}blip')
                        if blip:
                            rel_id = blip[0].get(f'{{{R_NS}}}embed')
                            if rel_id and rel_id in image_rels:
                                img_path = image_rels[rel_id]

                                if img_path not in media_cache:
                                    media_cache[img_path] = zip_ref.read(img_path)

                                img_data = media_cache[img_path]

                                ext = Path(img_path).suffix.lower().replace('.', '')
                                if ext in ('jpg', 'jpeg'):
                                    ext = 'jpeg'

                                import base64
                                img_b64 = base64.b64encode(img_data).decode('utf-8')

                                paragraph_data["runs"].append({
                                    "type": "image",
                                    "image_data": img_b64,
                                    "image_format": ext
                                })

                    if run_data := self._build_text_run(run_text, run_bold, run_italic, run_underline, run_font_name, run_font_size):
                        paragraph_data["runs"].append(run_data)

                self.ast["paragraphs"].append(paragraph_data)

    @staticmethod
    def _build_text_run(text, bold, italic, underline, font_name, font_size):
        if not text:
            return None
        run_data = {
            "text": text,
            "bold": bold,
            "italic": italic,
            "underline": underline,
            "font_name": font_name,
            "font_size": font_size
        }
        return run_data

    def _build_image_relationship_map(self):
        image_rels = {}
        try:
            with zipfile.ZipFile(self.docx_path, "r") as zip_ref:
                if "word/_rels/document.xml.rels" in zip_ref.namelist():
                    rels_xml = zip_ref.read("word/_rels/document.xml.rels")
                    rels_root = etree.fromstring(rels_xml)
                    for rel in rels_root:
                        rel_id = rel.get('Id')
                        target = rel.get('Target', '')
                        if target.startswith('media/'):
                            image_rels[rel_id] = f'word/{target}'
        except:
            pass
        return image_rels

    def parse_tables(self):
        for table_idx, table in enumerate(self.document.tables):

            table_data = {
                "id": table_idx,
                "rows": []
            }

            for row in table.rows:
                row_data = []

                for cell in row.cells:
                    row_data.append(cell.text)

                table_data["rows"].append(row_data)

            self.ast["tables"].append(table_data)

    def extract_images(self):
        with zipfile.ZipFile(self.docx_path, "r") as zip_ref:

            media_files = [
                f for f in zip_ref.namelist()
                if f.startswith("word/media/")
            ]

            for idx, image_path in enumerate(media_files):

                self.ast["images"].append({
                    "id": idx,
                    "path": image_path
                })

    def parse_formulas(self):

        with zipfile.ZipFile(self.docx_path, "r") as zip_ref:
            xml_content = zip_ref.read("word/document.xml")

        root = etree.fromstring(xml_content)

        namespaces = {
            "m": (
                "http://schemas.openxmlformats.org/"
                "officeDocument/2006/math"
            )
        }

        formulas = root.xpath(
            "//m:oMath",
            namespaces=namespaces
        )

        for idx, formula in enumerate(formulas):

            self.ast["formulas"].append({
                "id": idx,
                "xml": etree.tostring(
                    formula
                ).decode("utf-8")
            })

    def export_ast(self, output_path):

        output_path = Path(output_path)

        output_path.parent.mkdir(
            parents=True,
            exist_ok=True
        )

        with open(
            output_path,
            "w",
            encoding="utf-8"
        ) as f:

            json.dump(
                self.ast,
                f,
                ensure_ascii=False,
                indent=2
            )

    def run(self):

        self.validate_docx()

        self.parse_metadata()

        self.parse_styles()

        self.parse_paragraphs()

        self.parse_tables()

        self.extract_images()

        self.parse_formulas()

        self._identify_sections()

        self.export_ast(
            "workspace/parsed/document_ast.json"
        )

        print(
            "[INFO] AST exported successfully"
        )

    def _identify_sections(self):
        import re

        paras = self.ast.get("paragraphs", [])
        if not paras:
            return

        CHINESE_NUMS = "一二三四五六七八九十百"
        H1_PATTERN = re.compile(r'^第[{}0-9一二三四五六七八九十]+[章节部篇]'.format(CHINESE_NUMS))
        H1_PATTERN_NUM = re.compile(r'^\d+\s+\S')
        H2_PATTERN_CN = re.compile(r'^[一二三四五六七八九十]+[、．.]')
        H2_PATTERN_KW = re.compile(r'^(引言|摘要|结论|参考文献|致谢|附录|abstract|references|introduction|conclusion|summary)', re.IGNORECASE)
        H3_PATTERN = re.compile(r'^\d+\.\d+[\.\s]')

        # 封面页特征关键词
        COVER_KEYWORDS = [
            "摘要", "abstract", "关键词", "keywords", "key words",
            "作者", "author", "单位", "university", "学校", "学院",
            "指导教师", "supervisor", "论文", "thesis", "学位", "degree"
        ]

        def detect_heading_level(p, in_cover=False):
            """检测段落的标题级别
            
            严格的标题检测规则：
            1. 优先使用 Word 样式
            2. 模式匹配需要满足格式要求
            3. 避免将普通段落误识别为标题
            """
            style = (p.get("style") or "").lower()
            text = (p.get("text") or "").strip()
            if not text:
                return None, None
            
            # 优先使用 Word 样式（最可靠）
            if "heading 1" in style:
                return "Heading 1", text
            if "heading 2" in style:
                return "Heading 2", text
            if "heading 3" in style:
                return "Heading 3", text

            # 文本长度限制（标题通常较短）
            if len(text) > 80:
                return None, None

            # H1 检测：第X章/节/部分
            if H1_PATTERN.match(text):
                return "Heading 1", text
            
            # H1 检测：数字 + 空格 + 文本（如 "1  需求分析"）
            if H1_PATTERN_NUM.match(text) and len(text) < 50:
                return "Heading 1", text
            
            # H2 检测：中文数字（一、二、三）
            if H2_PATTERN_CN.match(text):
                return "Heading 2", text
            
            # H2 检测：关键词+冒号（引言：、摘要：）
            # 需要严格匹配：关键词后必须有冒号，且文本较短
            if H2_PATTERN_KW.match(text.lower()):
                # 检查是否有冒号
                if "：" in text or ":" in text:
                    # 如果在封面区域内，摘要/关键词不识别为标题
                    if in_cover:
                        return None, text
                    return "Heading 2", text
            
            # H3 检测：数字编号（1. xxx、1.1 xxx）
            if H3_PATTERN.match(text):
                # 对于 First Paragraph 样式，需要额外检查是否是标题
                if "first paragraph" in style:
                    # 检查是否是标题格式（如 "1.1 反应方程式"）
                    # 如果是数字编号格式，且长度较短，则识别为标题
                    if len(text) < 50:
                        return "Heading 3", text
                else:
                    return "Heading 3", text

            return None, None

        def is_cover_paragraph(p, next_p=None, prev_p=None):
            """检测段落是否可能是封面内容"""
            text = (p.get("text") or "").strip()
            style = (p.get("style") or "").lower()
            if not text:
                return True  # 空行视为封面内容
            
            # 包含页码标记
            if re.match(r'^---\s*PAGE\s+\d+\s*---$', text):
                return True
            
            # 包含封面关键词
            text_lower = text.lower()
            for kw in COVER_KEYWORDS:
                if kw in text_lower:
                    return True
            
            # 包含邮箱格式
            if "@" in text and ("." in text):
                return True
            
            # 论文标题特征：Heading 1 样式且较短，且后续有作者/单位信息
            if "heading 1" in style and len(text) < 50:
                # 检查后续段落是否包含作者/单位信息
                if next_p:
                    next_text = (next_p.get("text") or "").strip()
                    next_style = (next_p.get("style") or "").lower()
                    # 后续段落包含作者信息或单位信息
                    if any(kw in next_text.lower() for kw in ["作者", "author", "单位", "university", "学校", "学院"]):
                        return True
                    # 后续段落包含摘要/关键词
                    if any(kw in next_text.lower() for kw in ["摘要", "abstract", "关键词", "keywords"]):
                        return True
                    # 后续段落是普通段落且较短（可能是作者信息）
                    if not next_style.startswith("heading") and len(next_text) < 100:
                        return True
            
            # 作者信息特征：包含姓名和单位
            if re.match(r'^[A-Z][a-z]+\s+[A-Z][a-z]+', text) and "(" in text:
                return True
            if re.match(r'^[\u4e00-\u9fa5]+,\s*[\u4e00-\u9fa5]+', text) and "(" in text:
                return True
            
            return False

        cover_end = 0
        toc_start = -1
        toc_end = -1

        toc_keywords = ["目录", "目 录", "目  录", "table of contents", "contents"]

        # 第一遍扫描：检测目录区域
        for i, p in enumerate(paras):
            text = (p.get("text") or "").strip()
            text_lower = text.lower()
            if any(kw in text_lower for kw in toc_keywords):
                toc_start = i
            if "toc" in (p.get("style") or "").lower() or "目录" in (p.get("style") or "").lower():
                if toc_start < 0:
                    toc_start = i
                toc_end = i + 1

        # 第二遍扫描：检测封面区域
        # 策略：封面是第一个标题之前的所有段落，但需要识别包含摘要/关键词的完整封面
        body_start = -1
        
        # 首先找到第一个 Heading 1，作为可能的封面标题
        first_h1_idx = -1
        for i, p in enumerate(paras):
            text = (p.get("text") or "").strip()
            style = (p.get("style") or "").lower()
            if "heading 1" in style:
                first_h1_idx = i
                break
        
        # 如果找到第一个 Heading 1，检查它是否是封面标题
        if first_h1_idx >= 0:
            # 检查后续段落，确定封面范围
            cover_end_candidate = first_h1_idx + 1  # 至少包含标题本身
            
            # 向后扫描，检查是否包含作者、摘要、关键词等封面内容
            for i in range(first_h1_idx + 1, len(paras)):
                text = (paras[i].get("text") or "").strip()
                style = (paras[i].get("style") or "").lower()
                
                # 遇到目录区域，停止
                if toc_start >= 0 and toc_start <= i < toc_end:
                    break
                
                # 遇到正文起始关键词，停止
                text_lower = text.lower()
                if text_lower in ["引言", "前言", "introduction", "preface"]:
                    break
                
                # 遇到另一个标题，停止
                level, _ = detect_heading_level(paras[i], in_cover=False)
                if level and level != "Heading 1":
                    break
                if level == "Heading 1" and i > first_h1_idx:
                    break
                
                # 检查是否是封面内容
                next_p = paras[i + 1] if i + 1 < len(paras) else None
                prev_p = paras[i - 1] if i > 0 else None
                if is_cover_paragraph(paras[i], next_p, prev_p):
                    cover_end_candidate = i + 1
                else:
                    # 如果不是封面内容，但可能是作者信息（普通段落）
                    if not style.startswith("heading") and len(text) < 100:
                        cover_end_candidate = i + 1
                    else:
                        break
            
            body_start = cover_end_candidate
        else:
            # 没有找到 Heading 1，使用原始逻辑
            for i, p in enumerate(paras):
                text = (p.get("text") or "").strip()
                level, _ = detect_heading_level(p, in_cover=False)
                
                # 遇到目录区域，跳过
                if toc_start >= 0 and toc_start <= i < toc_end:
                    continue
                
                # 检测正文开始的标志
                if level in ("Heading 1", "Heading 2", "Heading 3"):
                    body_start = i
                    break
                
                # 检查是否是正文起始关键词
                text_lower = text.lower()
                if text_lower in ["引言", "前言", "introduction", "preface"]:
                    body_start = i
                    break

        # 如果没有找到正文开始，则查找第一个明显不是封面的段落
        if body_start == -1:
            for i, p in enumerate(paras):
                text = (p.get("text") or "").strip()
                if not text:
                    continue
                level, _ = detect_heading_level(p, in_cover=False)
                if level:
                    body_start = i
                    break
                # 如果段落较长且不含封面关键词，可能是正文
                if len(text) > 100 and not any(is_cover_paragraph(paras[j]) for j in range(max(0, i-2), i+1)):
                    body_start = i
                    break

        cover_end = body_start if body_start > 0 else 0

        # 如果仍未找到封面，使用原始逻辑
        if cover_end == 0:
            for i, p in enumerate(paras):
                text = (p.get("text") or "").strip()
                if text and len(text) > 5:
                    level, _ = detect_heading_level(p, in_cover=False)
                    if level:
                        cover_end = i
                        break

        # 完善目录区域检测
        if toc_start >= 0 and toc_end < 0:
            for i in range(toc_start + 1, len(paras)):
                level, _ = detect_heading_level(paras[i], in_cover=False)
                text = (paras[i].get("text") or "").strip()
                if level == "Heading 1" or (text and not level and len(text) > 20):
                    toc_end = i
                    break
            if toc_end < 0:
                toc_end = min(toc_start + 20, len(paras))

        # 标记段落区域，并根据标题检测结果更新样式
        for i, p in enumerate(paras):
            if toc_start >= 0 and toc_start <= i < toc_end:
                p["section"] = "toc"
            elif i < cover_end:
                p["section"] = "cover"
            else:
                p["section"] = "body"
                # 对于正文区域，根据标题检测结果更新样式
                level, _ = detect_heading_level(p, in_cover=False)
                if level:
                    p["style"] = level

        self.ast["section_regions"] = {
            "cover_end": cover_end,
            "toc_start": toc_start,
            "toc_end": toc_end
        }


if __name__ == "__main__":
    import sys

    input_path = sys.argv[1] if len(sys.argv) > 1 else "workspace/input/input.docx"
    parser = DocxParser(input_path)
    parser.run()