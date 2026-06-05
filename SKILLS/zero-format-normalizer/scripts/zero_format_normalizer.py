import json
import re
import zipfile
import tempfile
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None


class ZeroFormatNormalizer:
    def __init__(self, source_docx_path, template_config_path=None, edited_config_path=None):
        self.source_docx_path = Path(source_docx_path)
        self.template_config = self._load_template_config(template_config_path)
        self.edited_config = self._load_edited_config(edited_config_path)
        self._merge_edited_config()
        self.doc = None
        self.ast = {
            "paragraphs": [],
            "tables": [],
            "images": []
        }
        self._image_counter = 0

    def _load_template_config(self, config_path):
        if config_path and Path(config_path).exists():
            with open(config_path, "r", encoding="utf-8") as f:
                return json.load(f)
        return self._get_default_config()

    def _load_edited_config(self, config_path):
        if config_path and Path(config_path).exists():
            with open(config_path, "r", encoding="utf-8") as f:
                config = json.load(f)
                print(f"[INFO] Loaded edited config: {config_path}")
                return config
        return {}

    def _merge_edited_config(self):
        if not self.edited_config:
            return

        for key, value in self.edited_config.items():
            if key in ["cover_preserved", "redesign_cover", "toc_preserved"]:
                continue
            if isinstance(value, dict) and key in self.template_config:
                self._deep_merge(self.template_config[key], value)
            else:
                self.template_config[key] = value

    def _deep_merge(self, base, override):
        for key, value in override.items():
            if isinstance(value, dict) and key in base and isinstance(base[key], dict):
                self._deep_merge(base[key], value)
            else:
                base[key] = value

    def _get_default_config(self):
        return {
            "fonts": {
                "chinese": {"family": "宋体", "size": 12},
                "english": {"family": "Times New Roman", "size": 12}
            },
            "heading": {
                "level1": {"font": "黑体", "size": 14, "bold": True, "alignment": "center", "spacing_before": 24, "spacing_after": 18},
                "level2": {"font": "黑体", "size": 12, "bold": True, "alignment": "left", "spacing_before": 18, "spacing_after": 12},
                "level3": {"font": "宋体", "size": 12, "bold": True, "alignment": "left", "spacing_before": 12, "spacing_after": 6}
            },
            "paragraph": {
                "alignment": "justify",
                "line_spacing": 1.5,
                "first_indent": 2,
                "paragraph_spacing": False
            },
            "toc": {
                "enabled": True,
                "title": "目  录",
                "title_font": "黑体",
                "title_size": 16,
                "max_level": 3
            },
            "header": {
                "enabled": True,
                "text": "",
                "font": "宋体",
                "size": 9,
                "alignment": "center",
                "separator_line": True
            },
            "footer": {
                "enabled": True,
                "page_number_format": "arabic",
                "font": "宋体",
                "size": 9,
                "alignment": "center"
            },
            "image": {
                "max_width_percent": 80
            }
        }

    def extract_content(self):
        source_doc = Document(str(self.source_docx_path))

        image_rels = self._build_image_relationship_map()

        for para in source_doc.paragraphs:
            para_data = {
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "runs": []
            }

            M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
            for child in para._element:
                tag_local = child.tag.split('}')[1] if '}' in child.tag else child.tag
                ns_uri = child.tag.split('}')[0].strip('{') if '}' in child.tag else ''

                if ns_uri == M_NS and tag_local in ('oMath', 'oMathPara'):
                    xml_str = etree.tostring(child).decode('utf-8')
                    math_texts = []
                    for mt in child.iter(f'{{{M_NS}}}t'):
                        if mt.text:
                            math_texts.append(mt.text)
                    para_data["runs"].append({
                        "type": "formula",
                        "formula_type": tag_local,
                        "xml": xml_str,
                        "text": ''.join(math_texts)
                    })
                    continue

                if child.tag == f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}r':
                    run_elem = child
                    run_text = ''
                    for t_elem in run_elem.findall(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}t'):
                        if t_elem.text:
                            run_text += t_elem.text

                    run_bold = None
                    run_italic = None
                    run_font_name = None
                    run_font_size = None
                    rpr = run_elem.find(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}rPr')
                    if rpr is not None:
                        if rpr.find(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}b') is not None:
                            run_bold = True
                        if rpr.find(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}i') is not None:
                            run_italic = True
                        rFonts = rpr.find(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}rFonts')
                        if rFonts is not None:
                            run_font_name = rFonts.get(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}eastAsia') or rFonts.get(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}ascii')
                        sz = rpr.find(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}sz')
                        if sz is not None:
                            val = sz.get(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}val')
                            if val:
                                run_font_size = str(int(val) // 2) + 'pt'

                    run_data = {
                        "text": run_text,
                        "bold": run_bold,
                        "italic": run_italic,
                        "font_name": run_font_name,
                        "font_size": run_font_size
                    }
                    if run_text:
                        para_data["runs"].append(run_data)

            self.ast["paragraphs"].append(para_data)

        for table in source_doc.tables:
            table_data = {"rows": []}
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data["rows"].append(row_data)
            self.ast["tables"].append(table_data)

        self._extract_images_from_docx(image_rels)

    def _build_image_relationship_map(self):
        image_rels = {}
        try:
            with zipfile.ZipFile(self.source_docx_path, "r") as zip_ref:
                if "word/_rels/document.xml.rels" in zip_ref.namelist():
                    rels_xml = zip_ref.read("word/_rels/document.xml.rels")
                    from lxml import etree
                    rels_root = etree.fromstring(rels_xml)
                    for rel in rels_root:
                        rel_id = rel.get('Id')
                        target = rel.get('Target', '')
                        if target.startswith('media/'):
                            image_rels[rel_id] = f'word/{target}'
        except Exception as e:
            print(f"[WARNING] Failed to build image relationship map: {e}")
        return image_rels

    def _extract_images_from_docx(self, image_rels):
        try:
            with zipfile.ZipFile(self.source_docx_path, "r") as zip_ref:
                media_files = [f for f in zip_ref.namelist() if f.startswith("word/media/")]
                for idx, image_path in enumerate(media_files):
                    self.ast["images"].append({
                        "id": idx,
                        "path": image_path,
                        "data": zip_ref.read(image_path)
                    })
        except Exception as e:
            print(f"[WARNING] Failed to extract images: {e}")

    def _detect_structure(self):
        paras = self.ast["paragraphs"]
        cover_end = 0
        toc_start = -1
        toc_end = -1

        toc_keywords = ["目录", "目 录", "目  录", "table of contents", "contents"]

        first_heading_idx = -1
        for i, p in enumerate(paras):
            text = (p.get("text") or "").strip()
            text_lower = text.lower()

            if any(kw in text_lower for kw in toc_keywords):
                toc_start = i

            if self._detect_heading_level(text):
                if first_heading_idx < 0:
                    first_heading_idx = i
                if toc_start >= 0 and toc_end < 0:
                    toc_end = i
                    break

        if first_heading_idx > 0:
            cover_end = first_heading_idx
        elif first_heading_idx == 0:
            cover_end = 0

        if cover_end == 0 and first_heading_idx < 0:
            for i, p in enumerate(paras):
                text = (p.get("text") or "").strip()
                if text and len(text) > 5:
                    cover_end = i
                    break

        if toc_start >= 0 and toc_end < 0:
            for i in range(toc_start + 1, len(paras)):
                text = (paras[i].get("text") or "").strip()
                if self._detect_heading_level(text):
                    toc_end = i
                    break
            if toc_end < 0:
                toc_end = min(toc_start + 20, len(paras))

        for i, p in enumerate(paras):
            if toc_start >= 0 and toc_start <= i < toc_end:
                p["section"] = "toc"
            elif i < cover_end:
                p["section"] = "cover"
            else:
                p["section"] = "body"

        self.ast["section_regions"] = {
            "cover_end": cover_end,
            "toc_start": toc_start,
            "toc_end": toc_end
        }

    def _detect_heading_level(self, text, prev_text=None, next_text=None):
        text = text.strip()
        if not text:
            return None

        if re.match(r'^\d+\.\d+\.\d+\s+\S', text):
            return 3
        elif re.match(r'^\d+\.\d+\s+\S', text):
            return 2
        elif re.match(r'^\d+[\.\、]\s*\S', text) and len(text) <= 50:
            return 3

        if re.match(r'^[一二三四五六七八九十]+[\、\.\s]', text):
            return 2

        if re.match(r'^第[一二三四五六七八九十\d]+[章节部分篇]', text):
            return 1

        if re.match(r'^第[一二三四五六七八九十\d]+核心', text):
            return 2
        if re.match(r'^趋势[一二三四五六七八九十\d]', text):
            return 2

        if text.startswith('远景战略') or text.startswith('构筑') or text.startswith('面向'):
            if len(text) <= 60:
                return 1

        if '：' in text:
            parts = text.split('：', 1)
            prefix = parts[0].strip()
            if len(prefix) <= 15 and len(text) <= 80:
                if any(kw in prefix for kw in ['引言', '摘要', '结论', '总结', '概述', '前言', '背景']):
                    return 2
                if re.match(r'^[一二三四五六七八九十\d]+$', prefix):
                    return 2
                if len(prefix) <= 10:
                    return 2

        if prev_text is not None and next_text is not None:
            prev_len = len(prev_text.strip())
            next_len = len(next_text.strip())
            curr_len = len(text)
            if curr_len <= 40 and not any(c in text for c in '。，；、（）《》""'):
                if prev_len > 100 and next_len > 100:
                    return 2
                if prev_len == 0 and next_len > 80:
                    return 2
                if prev_len <= 40 and next_len > 120 and curr_len <= 25:
                    return 2

        return None

    def create_formatted_document(self):
        self.doc = Document()

        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

        page_cfg = self.template_config.get("page", {})
        section.top_margin = Cm(page_cfg.get("margin_top", 2.54))
        section.bottom_margin = Cm(page_cfg.get("margin_bottom", 2.54))
        section.left_margin = Cm(page_cfg.get("margin_left", 3.17))
        section.right_margin = Cm(page_cfg.get("margin_right", 2.54))

        self._detect_structure()

        self._add_cover_page()
        self._add_body_content()
        self._add_references_section()
        self._add_toc_page()
        self._create_sections()
        self._add_header_footer()
        self._enable_auto_update_fields()

    def _insert_cover_logo(self, logo_cfg):
        import base64
        import io
        
        img_data = logo_cfg.get("image_data", "")
        img_path = logo_cfg.get("image_path", "")
        
        img_stream = None
        
        if img_data and len(img_data) > 100:
            try:
                if img_data.startswith("data:"):
                    img_data = img_data.split(",", 1)[1]
                img_bytes = base64.b64decode(img_data)
                img_stream = io.BytesIO(img_bytes)
            except Exception as e:
                print(f"[WARNING] Failed to decode logo base64: {e}")
        
        if not img_stream and img_path:
            img_path_obj = Path(img_path)
            if img_path_obj.exists():
                try:
                    img_stream = open(img_path_obj, 'rb')
                except Exception as e:
                    print(f"[WARNING] Failed to open logo file: {e}")
        
        if img_stream:
            try:
                logo_para = self.doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_para.add_run()
                logo_width = Cm(logo_cfg.get("width", 120) / 50)
                logo_run.add_picture(img_stream, width=logo_width)
            except Exception as e:
                print(f"[WARNING] Failed to insert logo: {e}")
            finally:
                if hasattr(img_stream, 'close'):
                    img_stream.close()

    def _add_cover_page(self):
        cover_cfg = self.template_config.get("cover", {})
        if not cover_cfg.get("enabled", False):
            return

        logo_cfg = cover_cfg.get("logo", {})
        if logo_cfg.get("enabled", False):
            self._insert_cover_logo(logo_cfg)

        school_name = cover_cfg.get("school_name", "")
        if school_name:
            school_font = cover_cfg.get("school_font", "宋体")
            school_size = cover_cfg.get("school_size", 18)
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(school_name)
            run.bold = True
            run.font.size = Pt(school_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_run_font(run, school_font, True)
            self.doc.add_paragraph()

        layout_cfg = cover_cfg.get("layout", {})
        vertical_align = layout_cfg.get("vertical_align", "center")
        if vertical_align == "center":
            for _ in range(4):
                self.doc.add_paragraph()

        title_cfg = cover_cfg.get("title", {})
        title_text = title_cfg.get("text", "课程作业")
        title_font = title_cfg.get("font", "黑体")
        title_size = title_cfg.get("size", 22)
        title_bold = title_cfg.get("bold", True)
        title_align = title_cfg.get("alignment", "center")

        para = self.doc.add_paragraph()
        run = para.add_run(title_text)
        run.bold = title_bold
        run.font.size = Pt(title_size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        self._set_run_font(run, title_font, True)

        if title_align == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif title_align == "left":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif title_align == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for _ in range(4):
            self.doc.add_paragraph()

        info_items = cover_cfg.get("info_items", [])
        for item in info_items:
            label = item.get("label", "")
            value = item.get("value", "")
            if not value:
                continue
            item_font = item.get("font", "宋体")
            item_size = item.get("size", 14)

            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = para.add_run(f"{label}：{value}")
            run.font.size = Pt(item_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_run_font(run, item_font, True)

        self.doc.add_paragraph()

        self._add_section_break(is_cover=True)

    def _add_section_break(self, is_cover=False, is_toc=False, is_ref=False):
        last_para = self.doc.paragraphs[-1]._element
        ppr = last_para.find(qn('w:pPr'))
        if ppr is None:
            ppr = OxmlElement('w:pPr')
            last_para.insert(0, ppr)
        sect_pr = OxmlElement('w:sectPr')
        self._set_section_properties(sect_pr, is_cover=is_cover, is_toc=is_toc, is_ref=is_ref)
        ppr.append(sect_pr)

    def _add_toc_page(self):
        toc_cfg = self.template_config.get("toc", {})
        if not toc_cfg.get("enabled", False):
            return

        title = toc_cfg.get("title", "目  录")
        title_font = toc_cfg.get("title_font", "黑体")
        title_size = toc_cfg.get("title_size", 16)
        max_level = toc_cfg.get("max_level", 3)

        abstract_keywords = ["摘要", "摘 要", "abstract", "提要"]
        insert_target = None
        doc_paras = self.doc.paragraphs

        for i, para in enumerate(doc_paras):
            style_name = para.style.name if para.style else ""
            text = (para.text or "").strip().lower()
            if "heading" in style_name.lower():
                for kw in abstract_keywords:
                    if kw in text:
                        insert_target = para._element
                        break
                if insert_target is not None:
                    break

        if insert_target is None:
            for i, para in enumerate(doc_paras):
                style_name = para.style.name if para.style else ""
                if "heading" in style_name.lower():
                    insert_target = para._element
                    break

        if insert_target is None:
            doc_body = self.doc.element.body
            insert_target = doc_body

        title_elem = OxmlElement('w:p')
        title_ppr = OxmlElement('w:pPr')
        title_jc = OxmlElement('w:jc')
        title_jc.set(qn('w:val'), 'center')
        title_ppr.append(title_jc)
        title_spacing = OxmlElement('w:spacing')
        title_spacing.set(qn('w:before'), '240')
        title_spacing.set(qn('w:after'), '400')
        title_ppr.append(title_spacing)
        title_elem.append(title_ppr)
        title_r = OxmlElement('w:r')
        title_rpr = OxmlElement('w:rPr')
        title_b = OxmlElement('w:b')
        title_rpr.append(title_b)
        title_sz = OxmlElement('w:sz')
        title_sz.set(qn('w:val'), str(title_size * 2))
        title_rpr.append(title_sz)
        title_rfonts = OxmlElement('w:rFonts')
        title_rfonts.set(qn('w:ascii'), title_font)
        title_rfonts.set(qn('w:hAnsi'), title_font)
        title_rfonts.set(qn('w:eastAsia'), title_font)
        title_rpr.insert(0, title_rfonts)
        title_color = OxmlElement('w:color')
        title_color.set(qn('w:val'), '000000')
        title_rpr.append(title_color)
        title_r.append(title_rpr)
        title_t = OxmlElement('w:t')
        title_t.set(qn('xml:space'), 'preserve')
        title_t.text = title
        title_r.append(title_t)
        title_elem.append(title_r)

        toc_field_elem = OxmlElement('w:p')
        toc_ppr = OxmlElement('w:pPr')
        toc_field_elem.append(toc_ppr)

        r_begin = OxmlElement('w:r')
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        r_begin.append(fld_begin)
        toc_field_elem.append(r_begin)

        r_instr = OxmlElement('w:r')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = f' TOC \\o "1-{max_level}" \\h \\z \\u '
        r_instr.append(instr)
        toc_field_elem.append(r_instr)

        r_sep = OxmlElement('w:r')
        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        r_sep.append(fld_sep)
        toc_field_elem.append(r_sep)

        r_ph = OxmlElement('w:r')
        rph_rpr = OxmlElement('w:rPr')
        rph_color = OxmlElement('w:color')
        rph_color.set(qn('w:val'), '808080')
        rph_rpr.append(rph_color)
        r_ph.append(rph_rpr)
        t_ph = OxmlElement('w:t')
        t_ph.set(qn('xml:space'), 'preserve')
        t_ph.text = '（目录将在 Word 中自动更新）'
        r_ph.append(t_ph)
        toc_field_elem.append(r_ph)

        r_end = OxmlElement('w:r')
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        r_end.append(fld_end)
        toc_field_elem.append(r_end)

        sect_break_elem = OxmlElement('w:p')
        sect_ppr = OxmlElement('w:pPr')
        sect_pr = OxmlElement('w:sectPr')
        self._set_section_properties(sect_pr, is_toc=True)
        sect_ppr.append(sect_pr)
        sect_break_elem.append(sect_ppr)

        if insert_target is self.doc.element.body:
            doc_body = self.doc.element.body
            doc_body.insert(0, sect_break_elem)
            doc_body.insert(0, toc_field_elem)
            doc_body.insert(0, title_elem)
        else:
            insert_target.addprevious(title_elem)
            insert_target.addprevious(toc_field_elem)
            insert_target.addprevious(sect_break_elem)

    def _add_body_content(self):
        body_paras = [p for p in self.ast["paragraphs"] if p.get("section") == "body"]

        for i, para_data in enumerate(body_paras):
            text = para_data.get("text", "").strip()
            has_formula = any(r.get("type") == "formula" for r in para_data.get("runs", []))
            if not text and not has_formula:
                continue

            prev_text = body_paras[i-1].get("text", "") if i > 0 else None
            next_text = body_paras[i+1].get("text", "") if i < len(body_paras) - 1 else None

            heading_level = self._detect_heading_level(text, prev_text, next_text)

            if heading_level:
                self._add_heading(text, heading_level, para_data)
            else:
                self._add_normal_paragraph(text, para_data)

    def _add_heading(self, text, level, para_data=None):
        heading_cfg = self.template_config.get("heading", {})
        level_key = f"level{level}"
        cfg = heading_cfg.get(level_key, {})

        font_name = cfg.get("font", "黑体")
        font_size = cfg.get("size", 14 if level == 1 else 12)
        bold = cfg.get("bold", True)
        alignment_str = cfg.get("alignment", "center" if level == 1 else "left")
        spacing_before = cfg.get("spacing_before", 24 if level == 1 else 18)
        spacing_after = cfg.get("spacing_after", 18 if level == 1 else 12)

        heading_style_name = f'Heading {level}'
        try:
            para = self.doc.add_paragraph(style=heading_style_name)
        except KeyError:
            para = self.doc.add_paragraph()

        runs = para_data.get("runs", []) if para_data else []
        if runs:
            for run_data in runs:
                if run_data.get("type") == "formula":
                    self._add_formula_to_paragraph(para, run_data)
                    continue
                run = para.add_run(run_data.get("text", ""))
                run.bold = bold
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0, 0, 0)
                self._set_run_font(run, font_name, True)
        else:
            run = para.add_run(text)
            run.bold = bold
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_run_font(run, font_name, True)

        if alignment_str.upper() == "CENTER":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment_str.upper() == "LEFT":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment_str.upper() == "RIGHT":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if spacing_before:
            para.paragraph_format.space_before = Pt(spacing_before)
        if spacing_after:
            para.paragraph_format.space_after = Pt(spacing_after)

    def _add_normal_paragraph(self, text, para_data=None):
        para_cfg = self.template_config.get("paragraph", {})
        fonts_cfg = self.template_config.get("fonts", {})

        chinese_font = fonts_cfg.get("chinese", {}).get("family", "宋体")
        body_size = fonts_cfg.get("chinese", {}).get("size", 12)
        alignment = para_cfg.get("alignment", "justify").upper()
        line_spacing = para_cfg.get("line_spacing", 1.5)
        first_indent = para_cfg.get("first_indent", 2)

        para = self.doc.add_paragraph()

        runs = para_data.get("runs", []) if para_data else []
        if runs:
            for run_data in runs:
                if run_data.get("type") == "formula":
                    self._add_formula_to_paragraph(para, run_data)
                    continue
                run_text = run_data.get("text", "")
                run = para.add_run(run_text)
                is_chinese = any('\u4e00' <= c <= '\u9fff' for c in run_text)
                font_name = chinese_font if is_chinese else fonts_cfg.get("english", {}).get("family", "Times New Roman")
                self._set_run_font(run, font_name, is_chinese)
                run.font.size = Pt(body_size)
                if run_data.get("bold"):
                    run.bold = True
                if run_data.get("italic"):
                    run.italic = True
        else:
            run = para.add_run(text)
            is_chinese = any('\u4e00' <= c <= '\u9fff' for c in text)
            font_name = chinese_font if is_chinese else fonts_cfg.get("english", {}).get("family", "Times New Roman")
            self._set_run_font(run, font_name, is_chinese)
            run.font.size = Pt(body_size)

        if alignment == "JUSTIFY":
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif alignment == "CENTER":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "LEFT":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == "RIGHT":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        para.paragraph_format.line_spacing = line_spacing
        
        # 公式段落不添加首行缩进，且使用较小的段后间距
        has_formula = any(r.get("type") == "formula" for r in para_data.get("runs", []))
        if not has_formula:
            para.paragraph_format.first_line_indent = Pt(int(first_indent * body_size))
        else:
            # 公式段落使用较小的段后间距，避免空白过大
            para.paragraph_format.space_after = Pt(body_size * 0.5)

    def _add_references_section(self):
        ref_keywords = ["参考文献", "references", "bibliography"]
        ref_start = -1

        for i, p in enumerate(self.ast["paragraphs"]):
            text = (p.get("text") or "").strip().lower()
            if any(kw in text for kw in ref_keywords):
                ref_start = i
                break

        if ref_start < 0:
            return

        if self.doc.paragraphs:
            last_para = self.doc.paragraphs[-1]._element
            ppr = last_para.find(qn('w:pPr'))
            if ppr is None:
                ppr = OxmlElement('w:pPr')
                last_para.insert(0, ppr)
            sect_pr = OxmlElement('w:sectPr')
            self._set_section_properties(sect_pr, is_ref=True)
            ppr.append(sect_pr)

        para = self.doc.add_paragraph()
        ppr = para._element.get_or_add_pPr()
        pb = OxmlElement('w:pageBreakBefore')
        ppr.append(pb)

        run = para.add_run("参考文献")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
        self._set_run_font(run, "黑体", True)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(24)
        para.paragraph_format.space_after = Pt(18)

        for p in self.ast["paragraphs"][ref_start + 1:]:
            text = p.get("text", "").strip()
            if not text:
                continue

            para = self.doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(10.5)
            self._set_run_font(run, "宋体", True)
            para.paragraph_format.line_spacing = 1.5

    def _create_sections(self):
        doc_body = self.doc.element.body

        all_paras = self.doc.paragraphs
        if not all_paras:
            return

        last_para = all_paras[-1]._element
        ppr = last_para.find(qn('w:pPr'))
        if ppr is None:
            ppr = OxmlElement('w:pPr')
            last_para.insert(0, ppr)
        sect_pr = OxmlElement('w:sectPr')
        self._set_section_properties(sect_pr, is_final=True)
        ppr.append(sect_pr)

    def _set_section_properties(self, sect_pr, is_cover=False, is_toc=False, is_ref=False, is_final=False):
        pg_sz = OxmlElement('w:pgSz')
        pg_sz.set(qn('w:w'), '11906')
        pg_sz.set(qn('w:h'), '16838')
        sect_pr.append(pg_sz)

        pg_mar = OxmlElement('w:pgMar')
        pg_mar.set(qn('w:top'), '1440')
        pg_mar.set(qn('w:right'), '1800')
        pg_mar.set(qn('w:bottom'), '1440')
        pg_mar.set(qn('w:left'), '1800')
        pg_mar.set(qn('w:header'), '851')
        pg_mar.set(qn('w:footer'), '992')
        pg_mar.set(qn('w:gutter'), '0')
        sect_pr.append(pg_mar)

        if not is_cover and not is_toc:
            pg_num_type = OxmlElement('w:pgNumType')
            if not is_ref:
                pg_num_type.set(qn('w:start'), '1')
            sect_pr.append(pg_num_type)

    def _add_header_footer(self):
        header_cfg = self.template_config.get("header", {})
        footer_cfg = self.template_config.get("footer", {})
        
        # 检查是否有封面和目录
        has_cover = self._detect_cover_section()
        has_toc = self.template_config.get("toc", {}).get("enabled", False)
        
        # 计算正文开始的 section 索引
        body_section_start = 0
        if has_cover:
            body_section_start += 1
        if has_toc:
            body_section_start += 1

        sections = self.doc.sections
        if len(sections) < 1:
            return

        for i, section in enumerate(sections):
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False

            # 判断是否是正文 section
            is_body_section = (i >= body_section_start)
            
            if not is_body_section:
                # 封面和目录 section：设置空的页眉和页脚
                if header_cfg.get("enabled", False):
                    header = section.header
                    if header.paragraphs:
                        hp = header.paragraphs[0]
                        hp.clear()
                    else:
                        hp = header.add_paragraph()
                    hr = hp.add_run("")
                    hr.font.size = Pt(header_cfg.get("size", 9))
                if footer_cfg.get("enabled", False):
                    footer = section.footer
                    if footer.paragraphs:
                        fp = footer.paragraphs[0]
                        fp.clear()
                    else:
                        fp = footer.add_paragraph()
                    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # 正文 section：设置有内容的页眉和页脚
                if header_cfg.get("enabled", False):
                    header = section.header
                    if header.paragraphs:
                        hp = header.paragraphs[0]
                        hp.clear()
                    else:
                        hp = header.add_paragraph()

                    header_text = header_cfg.get("text", "")
                    header_font = header_cfg.get("font", "宋体")
                    header_size = header_cfg.get("size", 9)
                    header_align = header_cfg.get("alignment", "center")
                    header_sep = header_cfg.get("separator_line", True)

                    hr = hp.add_run(header_text)
                    hr.font.size = Pt(header_size)
                    self._set_run_font(hr, header_font, is_chinese=True)

                    if header_align == "center":
                        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif header_align == "left":
                        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif header_align == "right":
                        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    if header_sep:
                        ppr = hp._element.get_or_add_pPr()
                        pbdr = OxmlElement('w:pBdr')
                        bottom = OxmlElement('w:bottom')
                        bottom.set(qn('w:val'), 'single')
                        bottom.set(qn('w:sz'), '4')
                        bottom.set(qn('w:space'), '1')
                        bottom.set(qn('w:color'), '000000')
                        pbdr.append(bottom)
                        ppr.append(pbdr)

                if footer_cfg.get("enabled", False):
                    footer = section.footer
                    if footer.paragraphs:
                        fp = footer.paragraphs[0]
                        fp.clear()
                    else:
                        fp = footer.add_paragraph()

                    footer_font = footer_cfg.get("font", "宋体")
                    footer_size = footer_cfg.get("size", 9)
                    footer_align = footer_cfg.get("alignment", "center")

                    if footer_align == "center":
                        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif footer_align == "left":
                        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif footer_align == "right":
                        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    r_begin = OxmlElement('w:r')
                    fld_begin = OxmlElement('w:fldChar')
                    fld_begin.set(qn('w:fldCharType'), 'begin')
                    r_begin.append(fld_begin)
                    fp._element.append(r_begin)

                    r_instr = OxmlElement('w:r')
                    instr = OxmlElement('w:instrText')
                    instr.set(qn('xml:space'), 'preserve')
                    instr.text = ' PAGE '
                    r_instr.append(instr)
                    fp._element.append(r_instr)

                    r_end = OxmlElement('w:r')
                    fld_end = OxmlElement('w:fldChar')
                    fld_end.set(qn('w:fldCharType'), 'end')
                    r_end.append(fld_end)
                    fp._element.append(r_end)

                    for r in fp.runs:
                        r.font.size = Pt(footer_size)
                        self._set_run_font(r, footer_font, is_chinese=True)

    def _detect_cover_section(self):
        """检测是否有封面 section"""
        # 检查封面配置
        cover_cfg = self.template_config.get("cover", {})
        if cover_cfg.get("enabled", False):
            return True
        
        # 检查 AST 中是否有封面区域
        for p in self.ast.get("paragraphs", []):
            if p.get("section") == "cover":
                return True
        
        return False

    def _enable_auto_update_fields(self):
        settings_part = self.doc.settings.element
        update_fields = settings_part.find(qn('w:updateFields'))
        if update_fields is not None:
            settings_part.remove(update_fields)
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'false')
        settings_part.append(update_fields)

    @staticmethod
    def _add_formula_to_paragraph(para, run_data):
        formula_xml = run_data.get("xml", "")
        if not formula_xml:
            return
        try:
            p_elem = para._element
            formula_elem = etree.fromstring(formula_xml.encode('utf-8'))
            p_elem.append(formula_elem)
        except Exception as e:
            print(f"[WARNING] Formula insertion error: {e}")

    @staticmethod
    def _set_run_font(run, font_name, is_chinese=False):
        if not font_name:
            return
        run.font.name = font_name
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        if is_chinese:
            rFonts.set(qn('w:eastAsia'), font_name)

    def save(self, output_path):
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self._set_update_fields()
        self.doc.save(str(output_path))
        print(f"[INFO] Formatted DOCX saved to: {output_path}")

    def _set_update_fields(self):
        """Set updateFields flag so Word will auto-update TOC when opening"""
        from lxml import etree
        settings = self.doc.settings.element
        update_fields = settings.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}updateFields')
        if update_fields is None:
            update_fields = etree.SubElement(settings, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}updateFields')
        update_fields.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'true')
        print("[INFO] Set updateFields=true for auto TOC update")

    def run(self, output_path):
        print(f"[INFO] Loading source DOCX: {self.source_docx_path}")
        self.extract_content()

        print(f"[INFO] Extracted {len(self.ast['paragraphs'])} paragraphs")
        print(f"[INFO] Extracted {len(self.ast['tables'])} tables")
        print(f"[INFO] Extracted {len(self.ast['images'])} images")

        print("[INFO] Creating formatted document...")
        self.create_formatted_document()

        print("[INFO] Saving formatted document...")
        self.save(output_path)

        self._apply_table_processor(output_path)

        print("[INFO] Zero format normalization completed!")
        return {
            "paragraphs": len(self.ast['paragraphs']),
            "tables": len(self.ast['tables']),
            "images": len(self.ast['images']),
            "output_path": str(output_path)
        }

    def _apply_table_processor(self, output_path):
        try:
            import sys
            skills_dir = Path(__file__).parent.parent.parent
            sys.path.insert(0, str(skills_dir / 'table-processor' / 'scripts'))
            from table_processor import TableProcessor

            config_path = str(self.config_path) if hasattr(self, 'config_path') and self.config_path else None
            if config_path is None:
                config_path = str(Path(__file__).parent.parent.parent.parent / 'workspace' / 'validated' / 'template_config.json')

            processor = TableProcessor(str(output_path), config_path)
            result = processor.run()
            processor.save(str(output_path))
            print(f"[INFO] Table-processor: {result['tables_formatted']} tables, {result['captions_formatted']} captions formatted")
        except Exception as e:
            print(f"[WARN] Table-processor skipped: {e}")


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 3:
        print("Usage: python zero_format_normalizer.py <input.docx> <output.docx> [template.json] [edited_config.json]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]
    template_path = sys.argv[3] if len(sys.argv) > 3 else None
    edited_path = sys.argv[4] if len(sys.argv) > 4 else None

    normalizer = ZeroFormatNormalizer(input_path, template_path, edited_path)
    result = normalizer.run(output_path)
    print(f"[INFO] Result: {json.dumps(result, indent=2)}")
