import json
import base64
import shutil
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None


class ASTToDocxConverter:
    def __init__(self, ast_path, template_config_path=None, source_docx_path=None):
        self.ast_path = Path(ast_path)
        self.template_config_path = template_config_path
        self.source_docx_path = Path(source_docx_path) if source_docx_path else None
        self.ast = self.load_ast()
        self.template_config = self.load_template_config()
        self.doc = None
        self._image_counter = 0

    def load_ast(self):
        with open(self.ast_path, "r", encoding="utf-8") as f:
            return json.load(f)

    def load_template_config(self):
        if self.template_config_path and Path(self.template_config_path).exists():
            with open(self.template_config_path, "r", encoding="utf-8") as f:
                return json.load(f)
        return {}

    @staticmethod
    def set_run_font(run, font_name, is_chinese=False):
        if not font_name:
            return
        run.font.name = font_name
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        if is_chinese:
            rFonts.set(qn('w:eastAsia'), font_name)

    @staticmethod
    def get_heading_level(style_name):
        if not style_name:
            return None
        style_lower = style_name.lower()
        if 'heading 1' in style_lower or 'heading1' in style_lower:
            return 1
        elif 'heading 2' in style_lower or 'heading2' in style_lower:
            return 2
        elif 'heading 3' in style_lower or 'heading3' in style_lower:
            return 3
        return None

    @staticmethod
    def _validate_heading(para_data, heading_level):
        import re
        text = (para_data.get("text") or "").strip()
        if not text:
            return False

        if len(text) > 80:
            return False

        heading_pattern = r'^\d+(\.\d+)*\s*\S'
        if re.match(heading_pattern, text):
            return True

        if heading_level == 1:
            if len(text) > 30:
                return False
        elif heading_level == 2:
            if len(text) > 50:
                return False
        elif heading_level == 3:
            if len(text) > 60:
                return False

        ending_punctuations = ['。', '，', '；', '！', '？', '、', '.', ',', ';', '!', '?']
        if text and text[-1] in ending_punctuations:
            return False

        return True

    def setup_document(self):
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    def _get_available_width(self):
        try:
            section = self.doc.sections[0]
            page_width = section.page_width
            left_margin = section.left_margin
            right_margin = section.right_margin
            return page_width - left_margin - right_margin
        except (IndexError, AttributeError):
            return Cm(21) - Cm(3.18) - Cm(3.18)

    def _add_image_to_paragraph(self, para, run_data):
        img_data_b64 = run_data.get("image_data", "")
        img_format = run_data.get("image_format", "png")
        if not img_data_b64:
            return
        try:
            img_bytes = base64.b64decode(img_data_b64)
            suffix = '.png' if img_format == 'png' else '.jpg' if img_format in ('jpg', 'jpeg') else '.gif' if img_format == 'gif' else '.bmp' if img_format == 'bmp' else '.png'
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(img_bytes)
                tmp_path = tmp.name

            image_cfg = self.template_config.get("image", {})
            max_width_percent = image_cfg.get("max_width_percent", 80) / 100.0
            available_width = self._get_available_width()
            max_width = int(available_width * max_width_percent)

            img_width = None
            img_height = None
            if PILImage:
                try:
                    with PILImage.open(tmp_path) as pil_img:
                        img_width, img_height = pil_img.size
                except Exception:
                    pass

            run = para.add_run()
            if img_width and img_height and img_width > max_width:
                ratio = max_width / img_width
                new_height = int(img_height * ratio)
                run.add_picture(tmp_path, width=Emu(max_width), height=Emu(new_height))
            else:
                run.add_picture(tmp_path, width=Emu(max_width))

            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            Path(tmp_path).unlink(missing_ok=True)
            self._image_counter += 1
        except Exception as e:
            import traceback
            print(f"[WARNING] Image processing error: {e}")
            traceback.print_exc()

    @staticmethod
    def _clear_paragraph(para):
        p_elem = para._element
        for child in list(p_elem):
            tag_local = child.tag.split('}')[1] if '}' in child.tag else child.tag
            if tag_local == 'r':
                p_elem.remove(child)

    def _is_protected(self, para_data):
        section = para_data.get("section", "")
        style = (para_data.get("style") or "").strip()
        if section == "cover" and not style.startswith("Heading"):
            return False
        return section == "toc"

    @staticmethod
    def _normalize_alignment(alignment_str):
        if not alignment_str or alignment_str == "None":
            return None
        alignment_upper = alignment_str.upper()
        if "CENTER" in alignment_upper:
            return "CENTER"
        elif "RIGHT" in alignment_upper:
            return "RIGHT"
        elif "JUSTIFY" in alignment_upper:
            return "JUSTIFY"
        elif "LEFT" in alignment_upper:
            return "LEFT"
        return None

    def convert_paragraphs_in_place(self):
        doc_paras = self.doc.paragraphs
        ast_paras = self.ast.get("paragraphs", [])
        count = min(len(doc_paras), len(ast_paras))

        for i in range(count):
            para_data = ast_paras[i]

            if self._is_protected(para_data):
                continue

            if self._is_code_block(para_data):
                para = doc_paras[i]
                self._clear_paragraph(para)
                para._element.addnext(self._add_code_block_to_table(para_data)._element)
                continue

            runs = para_data.get("runs", [])
            all_image = runs and all(r.get("type") == "image" for r in runs)
            if not all_image:
                text = para_data.get("text", "")
                if not text.strip() and not any(r.get("type") == "image" for r in runs):
                    continue

            style_name = para_data.get("style", "")
            heading_level = self.get_heading_level(style_name)

            para = doc_paras[i]

            if heading_level is not None:
                if self._validate_heading(para_data, heading_level):
                    self._apply_heading(para_data, heading_level, para)
                else:
                    self._apply_normal_paragraph(para_data, para)
            else:
                self._apply_normal_paragraph(para_data, para)

    def _apply_heading(self, para_data, heading_level, para):
        heading_config = self.template_config.get("heading", {})
        level_key = f"level{heading_level}"
        lvl_cfg = heading_config.get(level_key, {})

        font_name = lvl_cfg.get("font", "黑体")
        font_size = lvl_cfg.get("size", 16 if heading_level == 1 else 14 if heading_level == 2 else 13)
        bold = lvl_cfg.get("bold", True)
        spacing_before = lvl_cfg.get("spacing_before", 24 if heading_level == 1 else 18 if heading_level == 2 else 12)
        spacing_after = lvl_cfg.get("spacing_after", 18 if heading_level == 1 else 12 if heading_level == 2 else 6)
        alignment_str = lvl_cfg.get("alignment", "center" if heading_level == 1 else "left")

        self._clear_paragraph(para)

        runs = para_data.get("runs", [])
        if runs:
            for run_data in runs:
                run = para.add_run(run_data.get("text", ""))
                run.bold = bold
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0, 0, 0)
                self.set_run_font(run, font_name, is_chinese=True)
        elif para_data.get("text"):
            run = para.add_run(para_data["text"])
            run.bold = bold
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            self.set_run_font(run, font_name, is_chinese=True)

        if alignment_str.upper() == "CENTER":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment_str.upper() == "LEFT":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment_str.upper() == "RIGHT":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if spacing_before:
            para.paragraph_format.space_before = Pt(spacing_before)
        if spacing_after:
            para.paragraph_format.space_after = Pt(spacing_after)

    def _get_body_font_size(self):
        fonts_cfg = self.template_config.get("fonts", {})
        return fonts_cfg.get("chinese", {}).get("size", 12)

    def _get_default_font(self, is_chinese=True):
        fonts_cfg = self.template_config.get("fonts", {})
        if is_chinese:
            return fonts_cfg.get("chinese", {}).get("family", "宋体")
        return fonts_cfg.get("english", {}).get("family", "Times New Roman")

    @staticmethod
    def _is_code_block(para_data):
        style = (para_data.get("style") or "").lower()
        if "code" in style or "source" in style:
            return True
        runs = para_data.get("runs", [])
        if not runs:
            return False
        code_fonts = ["consolas", "courier new", "monaco", "menlo", "dejavu sans mono", "liberation mono"]
        for run in runs:
            font = (run.get("font_name") or "").lower()
            if font in code_fonts:
                return True
        return False

    def _add_code_block_to_table(self, para_data):
        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        cell_para = cell.paragraphs[0]
        
        for run_data in para_data.get("runs", []):
            run = cell_para.add_run(run_data.get("text", ""))
            font_name = run_data.get("font_name") or "Consolas"
            run.font.name = font_name
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rpr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:eastAsia'), font_name)
            if run_data.get("font_size"):
                try:
                    size = int(run_data["font_size"].replace("pt", ""))
                    run.font.size = Pt(size)
                except Exception:
                    run.font.size = Pt(10)
            else:
                run.font.size = Pt(10)

        return table

    def _char_indent(self, char_count):
        body_size = self._get_body_font_size()
        return Pt(int(char_count * body_size))

    def _apply_normal_paragraph(self, para_data, para):
        runs = para_data.get("runs", [])
        has_image = any(r.get("type") == "image" for r in runs)

        self._clear_paragraph(para)

        para_cfg = self.template_config.get("paragraph", {})
        default_align = para_cfg.get("alignment", "justify").upper()

        if has_image:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            if default_align == "JUSTIFY":
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif default_align == "CENTER":
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif default_align == "LEFT":
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif default_align == "RIGHT":
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        first_indent = para_cfg.get("first_indent", 2)
        para.paragraph_format.first_line_indent = self._char_indent(first_indent)

        line_spacing = para_cfg.get("line_spacing", 1.5)
        para.paragraph_format.line_spacing = line_spacing

        if para_cfg.get("paragraph_spacing", False):
            body_size = self._get_body_font_size()
            para.paragraph_format.space_after = Pt(body_size)

        chinese_font = self._get_default_font(True)
        english_font = self._get_default_font(False)
        body_size = self._get_body_font_size()

        for run_data in runs:
            if run_data.get("type") == "image":
                self._add_image_to_paragraph(para, run_data)
                continue

            run = para.add_run(run_data.get("text", ""))
            if run_data.get("bold"):
                run.bold = True
            if run_data.get("italic"):
                run.italic = True

            text_content = run_data.get("text", "")
            is_chinese = any('\u4e00' <= c <= '\u9fff' for c in text_content)
            font_name = chinese_font if is_chinese else english_font
            self.set_run_font(run, font_name, is_chinese)
            run.font.size = Pt(body_size)

    def convert_paragraphs(self):
        para_cfg = self.template_config.get("paragraph", {})
        default_align = para_cfg.get("alignment", "justify").upper()
        chinese_font = self._get_default_font(True)
        english_font = self._get_default_font(False)
        body_size = self._get_body_font_size()

        for para_data in self.ast.get("paragraphs", []):
            if self._is_protected(para_data):
                continue

            if self._is_code_block(para_data):
                self._add_code_block_to_table(para_data)
                continue

            runs = para_data.get("runs", [])
            all_image_para = runs and all(r.get("type") == "image" for r in runs)

            if not all_image_para:
                text = para_data.get("text", "")
                if not text.strip() and not any(r.get("type") == "image" for r in runs):
                    continue

            style_name = para_data.get("style", "")
            heading_level = self.get_heading_level(style_name)

            if heading_level is not None:
                para = self._create_heading(para_data, heading_level)
            else:
                para = self.doc.add_paragraph()

            if heading_level is not None:
                pass
            else:
                has_image = any(r.get("type") == "image" for r in runs)
                if has_image:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    if default_align == "JUSTIFY":
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    elif default_align == "CENTER":
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif default_align == "LEFT":
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif default_align == "RIGHT":
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                first_indent = para_cfg.get("first_indent", 2)
                para.paragraph_format.first_line_indent = self._char_indent(first_indent)

                line_spacing = para_cfg.get("line_spacing", 1.5)
                para.paragraph_format.line_spacing = line_spacing

                if para_cfg.get("paragraph_spacing", False):
                    para.paragraph_format.space_after = Pt(body_size)

                for run_data in runs:
                    if run_data.get("type") == "image":
                        self._add_image_to_paragraph(para, run_data)
                        continue

                    run = para.add_run(run_data.get("text", ""))
                    if run_data.get("bold"):
                        run.bold = True
                    if run_data.get("italic"):
                        run.italic = True

                    text_content = run_data.get("text", "")
                    is_chinese = any('\u4e00' <= c <= '\u9fff' for c in text_content)
                    font_name = chinese_font if is_chinese else english_font
                    self.set_run_font(run, font_name, is_chinese)
                    run.font.size = Pt(body_size)

    def _create_heading(self, para_data, heading_level):
        heading_config = self.template_config.get("heading", {})
        level_key = f"level{heading_level}"
        lvl_cfg = heading_config.get(level_key, {})

        font_name = lvl_cfg.get("font", "黑体")
        font_size = lvl_cfg.get("size", 16 if heading_level == 1 else 14 if heading_level == 2 else 13)
        bold = lvl_cfg.get("bold", True)
        spacing_before = lvl_cfg.get("spacing_before", 24 if heading_level == 1 else 18 if heading_level == 2 else 12)
        spacing_after = lvl_cfg.get("spacing_after", 18 if heading_level == 1 else 12 if heading_level == 2 else 6)
        alignment_str = lvl_cfg.get("alignment", "center" if heading_level == 1 else "left")

        para = self.doc.add_paragraph()

        for run_data in para_data.get("runs", []):
            run = para.add_run(run_data.get("text", ""))
            run.bold = bold
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            self.set_run_font(run, font_name, is_chinese=True)

        if para_data.get("text") and not para_data.get("runs"):
            run = para.add_run(para_data["text"])
            run.bold = bold
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            self.set_run_font(run, font_name, is_chinese=True)

        if alignment_str.upper() == "CENTER":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment_str.upper() == "LEFT":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment_str.upper() == "RIGHT":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if spacing_before:
            para.paragraph_format.space_before = Pt(spacing_before)
        if spacing_after:
            para.paragraph_format.space_after = Pt(spacing_after)

        return para

    def convert_tables(self):
        for table_data in self.ast.get("tables", []):
            rows = table_data.get("rows", [])
            if not rows:
                continue

            table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.alignment = 1
            for i, row_data in enumerate(rows):
                for j, cell_text in enumerate(row_data):
                    cell = table.cell(i, j)
                    cell.text = cell_text
                    for para in cell.paragraphs:
                        for run in para.runs:
                            font_name = self._get_default_font(True)
                            run.font.name = font_name
                            rpr = run._element.get_or_add_rPr()
                            rFonts = rpr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = OxmlElement('w:rFonts')
                                rpr.insert(0, rFonts)
                            rFonts.set(qn('w:ascii'), font_name)
                            rFonts.set(qn('w:hAnsi'), font_name)
                            rFonts.set(qn('w:eastAsia'), font_name)
                            if not run.font.size:
                                run.font.size = Pt(self._get_body_font_size())

    def save(self, output_path):
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"[INFO] DOCX saved to: {output_path}")

    def _generate_toc_page(self):
        toc_cfg = self.template_config.get("toc", {})
        if not toc_cfg.get("enabled", False):
            return

        title = toc_cfg.get("title", "目  录")
        title_font = toc_cfg.get("title_font", "黑体")
        title_size = toc_cfg.get("title_size", 16)
        max_level = toc_cfg.get("max_level", 3)

        doc_paras = self.doc.paragraphs
        doc_body = self.doc.element.body

        cover_end = self._detect_real_cover_end()
        if cover_end > 0 and cover_end < len(doc_paras):
            insert_before = doc_paras[cover_end]._element
        elif len(doc_paras) > 0:
            insert_before = doc_paras[0]._element
        else:
            print("[WARNING] TOC generation: no paragraphs found, skipping")
            return

        title_elem = OxmlElement('w:p')
        title_ppr = OxmlElement('w:pPr')
        title_jc = OxmlElement('w:jc')
        title_jc.set(qn('w:val'), 'center')
        title_ppr.append(title_jc)
        title_spacing = OxmlElement('w:spacing')
        title_spacing.set(qn('w:before'), '240')
        title_spacing.set(qn('w:after'), '400')
        title_ppr.append(title_spacing)
        title_elem.append(title_ppr)
        title_r = OxmlElement('w:r')
        title_rpr = OxmlElement('w:rPr')
        title_b = OxmlElement('w:b')
        title_rpr.append(title_b)
        title_bcs = OxmlElement('w:bCs')
        title_rpr.append(title_bcs)
        title_sz = OxmlElement('w:sz')
        title_sz.set(qn('w:val'), str(title_size * 2))
        title_rpr.append(title_sz)
        title_szcs = OxmlElement('w:szCs')
        title_szcs.set(qn('w:val'), str(title_size * 2))
        title_rpr.append(title_szcs)
        title_rfonts = OxmlElement('w:rFonts')
        title_rfonts.set(qn('w:ascii'), title_font)
        title_rfonts.set(qn('w:hAnsi'), title_font)
        title_rfonts.set(qn('w:eastAsia'), title_font)
        title_rpr.insert(0, title_rfonts)
        title_color = OxmlElement('w:color')
        title_color.set(qn('w:val'), '000000')
        title_rpr.append(title_color)
        title_r.append(title_rpr)
        title_t = OxmlElement('w:t')
        title_t.set(qn('xml:space'), 'preserve')
        title_t.text = title
        title_r.append(title_t)
        title_elem.append(title_r)

        toc_field_elem = OxmlElement('w:p')
        toc_ppr = OxmlElement('w:pPr')
        toc_field_elem.append(toc_ppr)

        r_begin = OxmlElement('w:r')
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        r_begin.append(fld_begin)
        toc_field_elem.append(r_begin)

        r_instr = OxmlElement('w:r')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = f' TOC \\o "1-{max_level}" \\h \\z \\u '
        r_instr.append(instr)
        toc_field_elem.append(r_instr)

        r_sep = OxmlElement('w:r')
        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        r_sep.append(fld_sep)
        toc_field_elem.append(r_sep)

        r_ph = OxmlElement('w:r')
        rph_rpr = OxmlElement('w:rPr')
        rph_color = OxmlElement('w:color')
        rph_color.set(qn('w:val'), '808080')
        rph_rpr.append(rph_color)
        rph_sz = OxmlElement('w:sz')
        rph_sz.set(qn('w:val'), '21')
        rph_rpr.append(rph_sz)
        r_ph.append(rph_rpr)
        t_ph = OxmlElement('w:t')
        t_ph.set(qn('xml:space'), 'preserve')
        t_ph.text = '请在 Word 中按 Ctrl+A 然后按 F9 更新目录'
        r_ph.append(t_ph)
        toc_field_elem.append(r_ph)

        self._enable_auto_update_fields()

        r_end = OxmlElement('w:r')
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        r_end.append(fld_end)
        toc_field_elem.append(r_end)

        toc_section_break = OxmlElement('w:p')
        toc_sb_ppr = OxmlElement('w:pPr')
        toc_sb_sect = OxmlElement('w:sectPr')
        pg_sz = OxmlElement('w:pgSz')
        pg_sz.set(qn('w:w'), '11906')
        pg_sz.set(qn('w:h'), '16838')
        toc_sb_sect.append(pg_sz)
        pg_mar = OxmlElement('w:pgMar')
        pg_mar.set(qn('w:top'), '1440')
        pg_mar.set(qn('w:right'), '1800')
        pg_mar.set(qn('w:bottom'), '1440')
        pg_mar.set(qn('w:left'), '1800')
        pg_mar.set(qn('w:header'), '851')
        pg_mar.set(qn('w:footer'), '992')
        pg_mar.set(qn('w:gutter'), '0')
        toc_sb_sect.append(pg_mar)
        toc_sb_ppr.append(toc_sb_sect)
        toc_section_break.append(toc_sb_ppr)

        insert_before.addprevious(title_elem)
        insert_before.addprevious(toc_field_elem)
        insert_before.addprevious(toc_section_break)

        print(f"[INFO] TOC page generated (max_level={max_level})")

    def _enable_auto_update_fields(self):
        settings_part = self.doc.settings.element
        update_fields = settings_part.find(qn('w:updateFields'))
        if update_fields is None:
            update_fields = OxmlElement('w:updateFields')
            settings_part.append(update_fields)
        update_fields.set(qn('w:val'), 'true')

    def _apply_header_footer(self):
        header_cfg = self.template_config.get("header", {})
        footer_cfg = self.template_config.get("footer", {})
        toc_cfg = self.template_config.get("toc", {})

        has_cover = self._detect_real_cover_end() > 0
        has_toc = toc_cfg.get("enabled", False)
        pre_body_count = (1 if has_cover else 0) + (1 if has_toc else 0)

        sections = self.doc.sections
        if len(sections) < 1:
            return

        for i, section in enumerate(sections):
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False

            if i < pre_body_count:
                if header_cfg.get("enabled", False):
                    header = section.header
                    if header.paragraphs:
                        hp = header.paragraphs[0]
                        hp.clear()
                    else:
                        hp = header.add_paragraph()
                    hr = hp.add_run("")
                    hr.font.size = Pt(header_cfg.get("size", 9))
                if footer_cfg.get("enabled", False):
                    footer = section.footer
                    if footer.paragraphs:
                        fp = footer.paragraphs[0]
                        fp.clear()
                    else:
                        fp = footer.add_paragraph()
                    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:

                if header_cfg.get("enabled", False):
                    header = section.header
                    if header.paragraphs:
                        hp = header.paragraphs[0]
                        hp.clear()
                    else:
                        hp = header.add_paragraph()

                    header_text = header_cfg.get("text", "")
                    header_font = header_cfg.get("font", "宋体")
                    header_size = header_cfg.get("size", 9)
                    header_align = header_cfg.get("alignment", "center")
                    header_sep = header_cfg.get("separator_line", True)

                    hr = hp.add_run(header_text)
                    hr.font.size = Pt(header_size)
                    self.set_run_font(hr, header_font, is_chinese=True)

                    if header_align == "center":
                        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif header_align == "left":
                        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif header_align == "right":
                        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    if header_sep:
                        ppr = hp._element.get_or_add_pPr()
                        pbdr = OxmlElement('w:pBdr')
                        bottom = OxmlElement('w:bottom')
                        bottom.set(qn('w:val'), 'single')
                        bottom.set(qn('w:sz'), '4')
                        bottom.set(qn('w:space'), '1')
                        bottom.set(qn('w:color'), '000000')
                        pbdr.append(bottom)
                        ppr.append(pbdr)

                if footer_cfg.get("enabled", False):
                    footer = section.footer
                    if footer.paragraphs:
                        fp = footer.paragraphs[0]
                        fp.clear()
                    else:
                        fp = footer.add_paragraph()

                    footer_font = footer_cfg.get("font", "宋体")
                    footer_size = footer_cfg.get("size", 9)
                    footer_align = footer_cfg.get("alignment", "center")

                    if footer_align == "center":
                        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif footer_align == "left":
                        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif footer_align == "right":
                        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    from lxml import etree
                    r_begin = OxmlElement('w:r')
                    fld_begin = OxmlElement('w:fldChar')
                    fld_begin.set(qn('w:fldCharType'), 'begin')
                    r_begin.append(fld_begin)
                    fp._element.append(r_begin)

                    r_instr = OxmlElement('w:r')
                    instr = OxmlElement('w:instrText')
                    instr.set(qn('xml:space'), 'preserve')
                    instr.text = ' PAGE '
                    r_instr.append(instr)
                    fp._element.append(r_instr)

                    r_end = OxmlElement('w:r')
                    fld_end = OxmlElement('w:fldChar')
                    fld_end.set(qn('w:fldCharType'), 'end')
                    r_end.append(fld_end)
                    fp._element.append(r_end)

                    for r in fp.runs:
                        r.font.size = Pt(footer_size)
                        self.set_run_font(r, footer_font, is_chinese=True)

    def _fix_table_fonts_in_place(self):
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            font_name = self._get_default_font(True)
                            run.font.name = font_name
                            rpr = run._element.get_or_add_rPr()
                            rFonts = rpr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = OxmlElement('w:rFonts')
                                rpr.insert(0, rFonts)
                            rFonts.set(qn('w:ascii'), font_name)
                            rFonts.set(qn('w:hAnsi'), font_name)
                            rFonts.set(qn('w:eastAsia'), font_name)
                            if not run.font.size:
                                run.font.size = Pt(self._get_body_font_size())

    def _insert_page_break_before_references(self):
        ast_paras = self.ast.get("paragraphs", [])
        doc_paras = self.doc.paragraphs
        count = min(len(doc_paras), len(ast_paras))
        for i in range(count):
            para_data = ast_paras[i]
            text = (para_data.get("text") or "").strip()
            style = para_data.get("style", "") or ""
            if "Heading 1" in style and "参考文献" in text:
                para = doc_paras[i]
                p_elem = para._element
                ppr = p_elem.find(qn('w:pPr'))
                if ppr is None:
                    ppr = OxmlElement('w:pPr')
                    p_elem.insert(0, ppr)
                existing_pb = False
                for child in ppr:
                    tag_local = child.tag.split('}')[1] if '}' in child.tag else child.tag
                    if tag_local == 'pageBreakBefore':
                        existing_pb = True
                        break
                if not existing_pb:
                    pb = OxmlElement('w:pageBreakBefore')
                    ppr.append(pb)
                break

    def _remove_all_sections(self):
        doc_body = self.doc.element.body
        sect_prs = doc_body.findall(qn('w:sectPr'))
        for sect_pr in sect_prs:
            doc_body.remove(sect_pr)
        
        for para in self.doc.paragraphs:
            p_elem = para._element
            ppr = p_elem.find(qn('w:pPr'))
            if ppr is not None:
                sect_pr = ppr.find(qn('w:sectPr'))
                if sect_pr is not None:
                    ppr.remove(sect_pr)

    def _detect_real_cover_end(self):
        ast_paras = self.ast.get("paragraphs", [])
        if not ast_paras:
            return 0

        parser_cover_end = 0
        for i, p in enumerate(ast_paras):
            if p.get("section") == "cover":
                parser_cover_end = max(parser_cover_end, i + 1)

        if parser_cover_end == 0:
            return 0

        has_heading_in_cover = False
        for i in range(min(parser_cover_end, len(ast_paras))):
            style = (ast_paras[i].get("style") or "").strip()
            if style.startswith("Heading"):
                has_heading_in_cover = True
                break

        if has_heading_in_cover and parser_cover_end > 5:
            for i in range(min(parser_cover_end, len(ast_paras))):
                text = (ast_paras[i].get("text") or "").strip()
                if not text:
                    continue
                style = (ast_paras[i].get("style") or "").strip()
                if not style.startswith("Heading"):
                    has_heading_in_cover = False
                    break
            if has_heading_in_cover:
                print(f"[INFO] Cover section contains headings ({parser_cover_end} paras) - treating as body content")
                return 0

        return parser_cover_end

    def _normalize_sections(self):
        ast_paras = self.ast.get("paragraphs", [])
        doc_paras = self.doc.paragraphs
        count = min(len(doc_paras), len(ast_paras))
        doc_body = self.doc.element.body

        import re
        ref_pattern = re.compile(r'(参考文献|references|bibliography)', re.IGNORECASE)
        cover_end = self._detect_real_cover_end()

        ref_start = count
        for i, p in enumerate(ast_paras):
            text = (p.get("text") or "").strip()
            if ref_pattern.search(text):
                ref_start = i
                break

        has_cover = cover_end > 0 and cover_end < count

        if has_cover and cover_end < len(doc_paras):
            cover_last = doc_paras[cover_end - 1]._element
            ppr = cover_last.find(qn('w:pPr'))
            if ppr is None:
                ppr = OxmlElement('w:pPr')
                cover_last.insert(0, ppr)
            sect_pr = OxmlElement('w:sectPr')
            self._set_section_properties(sect_pr, is_cover=True)
            ppr.append(sect_pr)

        if ref_start < count and ref_start < len(doc_paras):
            ref_para = doc_paras[ref_start]._element
            ppr = ref_para.find(qn('w:pPr'))
            if ppr is None:
                ppr = OxmlElement('w:pPr')
                ref_para.insert(0, ppr)

            existing_pb = False
            for child in ppr:
                tag_local = child.tag.split('}')[1] if '}' in child.tag else child.tag
                if tag_local == 'pageBreakBefore':
                    existing_pb = True
                    break
            if not existing_pb:
                pb = OxmlElement('w:pageBreakBefore')
                ppr.append(pb)

            sect_pr = OxmlElement('w:sectPr')
            self._set_section_properties(sect_pr, is_ref=True)
            ppr.append(sect_pr)

        final_sect_pr = OxmlElement('w:sectPr')
        self._set_section_properties(final_sect_pr, is_final=True)
        doc_body.append(final_sect_pr)

    def _set_section_properties(self, sect_pr, is_cover=False, is_toc=False, is_ref=False, is_final=False):
        pg_sz = OxmlElement('w:pgSz')
        pg_sz.set(qn('w:w'), '11906')
        pg_sz.set(qn('w:h'), '16838')
        sect_pr.append(pg_sz)

        pg_mar = OxmlElement('w:pgMar')
        pg_mar.set(qn('w:top'), '1440')
        pg_mar.set(qn('w:right'), '1800')
        pg_mar.set(qn('w:bottom'), '1440')
        pg_mar.set(qn('w:left'), '1800')
        pg_mar.set(qn('w:header'), '851')
        pg_mar.set(qn('w:footer'), '992')
        pg_mar.set(qn('w:gutter'), '0')
        sect_pr.append(pg_mar)

        if not is_cover and not is_toc:
            pg_num_type = OxmlElement('w:pgNumType')
            if is_final:
                pg_num_type.set(qn('w:start'), '1')
            sect_pr.append(pg_num_type)

    def _load_edited_config(self):
        edited_config_path = Path("workspace/validated/edited_config.json")
        if edited_config_path.exists():
            with open(edited_config_path, "r", encoding="utf-8") as f:
                return json.load(f)
        return None

    def _detect_original_cover_range(self):
        ast_paras = self.ast.get("paragraphs", [])
        if not ast_paras:
            return (0, 0)

        parser_cover_end = 0
        for i, p in enumerate(ast_paras):
            if p.get("section") == "cover":
                parser_cover_end = max(parser_cover_end, i + 1)

        if parser_cover_end == 0:
            return (0, 0)

        has_heading_in_cover = False
        for i in range(parser_cover_end):
            style = (ast_paras[i].get("style") or "").strip()
            if style.startswith("Heading"):
                has_heading_in_cover = True
                break

        if has_heading_in_cover and parser_cover_end > 5:
            first_heading_idx = -1
            for i, p in enumerate(ast_paras):
                style = (p.get("style") or "").strip()
                if style.startswith("Heading"):
                    first_heading_idx = i
                    break
            if first_heading_idx > 0:
                return (0, first_heading_idx)
            return (0, 0)

        return (0, parser_cover_end)

    def _apply_cover_redesign(self, edited_config):
        cover_cfg = edited_config.get("cover", {})
        if not cover_cfg.get("enabled", False):
            return

        cover_start, cover_end = self._detect_original_cover_range()

        doc_paras = self.doc.paragraphs
        body = self.doc.element.body

        if cover_end > 0 and cover_end <= len(doc_paras):
            for i in range(cover_end - 1, cover_start - 1, -1):
                p_elem = doc_paras[i]._element
                body.remove(p_elem)
            self.ast["paragraphs"] = self.ast["paragraphs"][cover_end:]

        new_cover_paras = []

        logo_cfg = cover_cfg.get("logo", {})
        if logo_cfg.get("enabled", False):
            self._insert_cover_logo(logo_cfg, new_cover_paras)

        school_name = cover_cfg.get("school_name", "")
        if school_name:
            school_font = cover_cfg.get("school_font", "宋体")
            school_size = cover_cfg.get("school_size", 18)
            p = OxmlElement('w:p')
            ppr = OxmlElement('w:pPr')
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            ppr.append(jc)
            p.append(ppr)
            r = OxmlElement('w:r')
            rpr = OxmlElement('w:rPr')
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:ascii'), school_font)
            rf.set(qn('w:hAnsi'), school_font)
            rf.set(qn('w:eastAsia'), school_font)
            rpr.append(rf)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(school_size * 2)))
            rpr.append(sz)
            b = OxmlElement('w:b')
            rpr.append(b)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '000000')
            rpr.append(color)
            r.append(rpr)
            t = OxmlElement('w:t')
            t.text = school_name
            r.append(t)
            p.append(r)
            new_cover_paras.append(p)

            sp = OxmlElement('w:p')
            new_cover_paras.append(sp)

        layout_cfg = cover_cfg.get("layout", {})
        vertical_align = layout_cfg.get("vertical_align", "center")
        if vertical_align == "center":
            for _ in range(4):
                new_cover_paras.append(OxmlElement('w:p'))

        title_cfg = cover_cfg.get("title", {})
        title_text = title_cfg.get("text", "课程作业")
        title_font = title_cfg.get("font", "黑体")
        title_size = title_cfg.get("size", 22)
        title_bold = title_cfg.get("bold", True)
        title_align = title_cfg.get("alignment", "center")

        tp = OxmlElement('w:p')
        tppr = OxmlElement('w:pPr')
        tjc = OxmlElement('w:jc')
        tjc.set(qn('w:val'), title_align if title_align else 'center')
        tppr.append(tjc)
        tp.append(tppr)
        tr = OxmlElement('w:r')
        trpr = OxmlElement('w:rPr')
        trf = OxmlElement('w:rFonts')
        trf.set(qn('w:ascii'), title_font)
        trf.set(qn('w:hAnsi'), title_font)
        trf.set(qn('w:eastAsia'), title_font)
        trpr.append(trf)
        tsz = OxmlElement('w:sz')
        tsz.set(qn('w:val'), str(int(title_size * 2)))
        trpr.append(tsz)
        if title_bold:
            trpr.append(OxmlElement('w:b'))
        tcolor = OxmlElement('w:color')
        tcolor.set(qn('w:val'), '000000')
        trpr.append(tcolor)
        tr.append(trpr)
        tt = OxmlElement('w:t')
        tt.text = title_text
        tr.append(tt)
        tp.append(tr)
        new_cover_paras.append(tp)

        for _ in range(4):
            new_cover_paras.append(OxmlElement('w:p'))

        info_items = cover_cfg.get("info_items", [])
        for item in info_items:
            label = item.get("label", "")
            value = item.get("value", "")
            if not value:
                continue
            item_font = item.get("font", "宋体")
            item_size = item.get("size", 14)

            ip = OxmlElement('w:p')
            ippr = OxmlElement('w:pPr')
            ijc = OxmlElement('w:jc')
            ijc.set(qn('w:val'), 'center')
            ippr.append(ijc)
            ip.append(ippr)
            ir = OxmlElement('w:r')
            irpr = OxmlElement('w:rPr')
            irf = OxmlElement('w:rFonts')
            irf.set(qn('w:ascii'), item_font)
            irf.set(qn('w:hAnsi'), item_font)
            irf.set(qn('w:eastAsia'), item_font)
            irpr.append(irf)
            isz = OxmlElement('w:sz')
            isz.set(qn('w:val'), str(int(item_size * 2)))
            irpr.append(isz)
            icolor = OxmlElement('w:color')
            icolor.set(qn('w:val'), '000000')
            irpr.append(icolor)
            ir.append(irpr)
            it = OxmlElement('w:t')
            it.text = f"{label}：{value}"
            ir.append(it)
            ip.append(ir)
            new_cover_paras.append(ip)

        new_cover_paras.append(OxmlElement('w:p'))

        sect_break = OxmlElement('w:p')
        sect_ppr = OxmlElement('w:pPr')
        sect_pr = OxmlElement('w:sectPr')
        self._set_section_properties(sect_pr, is_cover=True)
        sect_ppr.append(sect_pr)
        sect_break.append(sect_ppr)
        new_cover_paras.append(sect_break)

        if doc_paras:
            first_elem = self.doc.paragraphs[0]._element
            body_elem = first_elem.getparent()
            insert_pos = list(body_elem).index(first_elem)
            for idx, p in enumerate(new_cover_paras):
                body_elem.insert(insert_pos + idx, p)
        else:
            for p in new_cover_paras:
                body.append(p)

        self.ast["paragraphs"] = [{
            "id": 0, "text": "", "style": "Normal", "section": "cover", "runs": []
        }] * len(new_cover_paras) + self.ast["paragraphs"]

        print(f"[INFO] Cover redesigned: title='{title_text}', {len(info_items)} info items")

    def _insert_cover_logo(self, logo_cfg, new_cover_paras):
        import base64
        import io

        img_data = logo_cfg.get("image_data", "")
        img_path = logo_cfg.get("image_path", "")

        if not img_data and not img_path:
            return

        try:
            if img_data and len(img_data) > 100:
                if img_data.startswith("data:"):
                    img_data = img_data.split(",", 1)[1]
                img_bytes = base64.b64decode(img_data)
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                    tmp.write(img_bytes)
                    tmp_path = tmp.name
            elif img_path and Path(img_path).exists():
                tmp_path = str(img_path)
            else:
                return

            logo_width_cm = logo_cfg.get("width", 120) / 50
            logo_para = OxmlElement('w:p')
            logo_ppr = OxmlElement('w:pPr')
            logo_jc = OxmlElement('w:jc')
            logo_jc.set(qn('w:val'), 'center')
            logo_ppr.append(logo_jc)
            logo_para.append(logo_ppr)
            new_cover_paras.append(logo_para)
            print(f"[INFO] Cover logo prepared")
        except Exception as e:
            print(f"[WARNING] Failed to prepare cover logo: {e}")

    def run(self, output_path):
        edited_config = self._load_edited_config()

        if self.source_docx_path and self.source_docx_path.exists():
            output = Path(output_path)
            output.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(self.source_docx_path, output)
            self.doc = Document(str(output))
            self._remove_all_sections()

            if edited_config and edited_config.get("redesign_cover", False):
                self._apply_cover_redesign(edited_config)
            else:
                self.ast["paragraphs"] = self.ast.get("paragraphs", [])

            self.convert_paragraphs_in_place()
            self._fix_table_fonts_in_place()
            self._normalize_sections()
            self._generate_toc_page()
            self._apply_header_footer()
            self.save(output_path)
        else:
            self.doc = Document()
            self.setup_document()
            self.convert_paragraphs()
            self.convert_tables()
            self._normalize_sections()
            self._generate_toc_page()
            self._apply_header_footer()
            self.save(output_path)


if __name__ == "__main__":
    converter = ASTToDocxConverter(
        "workspace/normalized/normalized_ast.json",
        "workspace/validated/template_config.json",
        "workspace/input/source.docx"
    )
    converter.run("workspace/output/final.docx")
