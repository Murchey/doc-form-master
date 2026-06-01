import re
import json
import copy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml


class TableProcessor:

    def __init__(self, docx_path, config_path=None):
        self.docx_path = Path(docx_path)
        self.doc = Document(str(self.docx_path))
        self.config = self._load_config(config_path)
        self.report = {
            "total_tables": 0,
            "total_captions": 0,
            "tables_formatted": 0,
            "captions_formatted": 0,
            "images_preserved": 0,
            "errors": []
        }

    def _load_config(self, config_path):
        if config_path and Path(config_path).exists():
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}

    def run(self):
        table_cfg = self.config.get("table", {})
        caption_cfg = self.config.get("caption", {}).get("table", {})
        body_cfg = self.config.get("body", {})
        heading_cfg = self.config.get("heading", {})

        caption_style = self._build_caption_style(caption_cfg, heading_cfg, body_cfg)
        header_style = self._build_header_style(table_cfg, heading_cfg)
        cell_style = self._build_cell_style(table_cfg, body_cfg)

        self._process_captions(caption_style)
        self._process_tables(table_cfg, header_style, cell_style)
        self._ensure_captions_with_tables()
        return self.report

    def _build_caption_style(self, caption_cfg, heading_cfg, body_cfg):
        font_family = self._get_font(heading_cfg, 2, "黑体")
        return {
            "font_name": font_family,
            "font_size": caption_cfg.get("font_size", 10.5),
            "bold": caption_cfg.get("bold", True),
            "alignment": self._parse_alignment(caption_cfg.get("alignment", "center")),
            "color": RGBColor(0, 0, 0)
        }

    def _build_header_style(self, table_cfg, heading_cfg):
        font_family = self._get_font(heading_cfg, 3, "黑体")
        return {
            "font_name": font_family,
            "font_size": table_cfg.get("header_font_size", 10.5),
            "bold": True,
            "alignment": WD_ALIGN_PARAGRAPH.CENTER
        }

    def _build_cell_style(self, table_cfg, body_cfg):
        font_family = body_cfg.get("font", {}).get("chinese", {}).get("primary", "宋体")
        return {
            "font_name": font_family,
            "font_size": table_cfg.get("cell_font_size", 10.5),
            "bold": False,
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "line_spacing": 1.0
        }

    def _get_font(self, heading_cfg, level, default):
        key = f"level{level}"
        level_cfg = heading_cfg.get(key, {})
        families = level_cfg.get("font_family", [])
        if isinstance(families, list) and families:
            return families[0]
        if isinstance(families, str):
            return families
        return default

    def _parse_alignment(self, align_str):
        mapping = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        return mapping.get(align_str, WD_ALIGN_PARAGRAPH.CENTER)

    # ---- Caption detection and formatting ----

    _TABLE_CAPTION_PATTERNS = [
        re.compile(r'^表\s*[-.\u2013]\s*\d+\s*[-.\u2013]\s*\d+'),
        re.compile(r'^表\s*\d+\s*[-.\u2013]\s*\d+'),
        re.compile(r'^表\s*\d+\s+'),
        re.compile(r'^Table\s*\d+\s*', re.IGNORECASE),
        re.compile(r'^表格\s*\d+\s+'),
    ]

    def _is_table_caption(self, text):
        text = text.strip()
        if not text:
            return False
        for pat in self._TABLE_CAPTION_PATTERNS:
            if pat.match(text):
                return True
        return False

    def _is_caption_style(self, para):
        style_name = para.style.name if para.style else ""
        return "caption" in style_name.lower()

    def _process_captions(self, style):
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            is_match = self._is_table_caption(text)
            if not is_match and self._is_caption_style(para):
                is_match = self._is_table_caption(text)

            if is_match:
                self._apply_caption_format(para, style)
                self.report["captions_formatted"] += 1
        self.report["total_captions"] = self.report["captions_formatted"]

    def _apply_caption_format(self, para, style):
        pf = para.paragraph_format
        pf.alignment = style["alignment"]
        pf.space_before = Pt(6)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.0

        for run in para.runs:
            run.font.name = style["font_name"]
            run.font.size = Pt(style["font_size"])
            run.font.bold = style["bold"]
            run.font.color.rgb = style["color"]
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rpr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), style["font_name"])
            rFonts.set(qn('w:hAnsi'), style["font_name"])
            rFonts.set(qn('w:eastAsia'), style["font_name"])

    # ---- Table formatting ----

    def _process_tables(self, table_cfg, header_style, cell_style):
        border_cfg = table_cfg.get("border", "three-line")
        cell_margin = table_cfg.get("cell_margin", 0.1)

        for table in self.doc.tables:
            self.report["total_tables"] += 1
            try:
                self._format_table_alignment(table)
                self._format_table_borders(table, border_cfg)
                self._format_table_autofit(table)
                self._format_table_cell_margins(table, cell_margin)
                self._format_header_row(table, header_style)
                self._format_body_rows(table, cell_style)
                self._count_images_in_table(table)
                self.report["tables_formatted"] += 1
            except Exception as e:
                self.report["errors"].append(str(e))

    def _format_table_alignment(self, table):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def _format_table_borders(self, table, border_cfg):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)

        if border_cfg == "three-line":
            borders_xml = self._build_three_line_borders()
        elif border_cfg == "grid":
            borders_xml = self._build_grid_borders()
        elif border_cfg == "none":
            borders_xml = self._build_no_borders()
        else:
            borders_xml = self._build_three_line_borders()

        tblPr.append(borders_xml)

    def _build_three_line_borders(self):
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'bottom']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            borders.append(border)

        inside_h = OxmlElement('w:insideH')
        inside_h.set(qn('w:val'), 'single')
        inside_h.set(qn('w:sz'), '6')
        inside_h.set(qn('w:space'), '0')
        inside_h.set(qn('w:color'), '000000')
        borders.append(inside_h)

        for border_name in ['left', 'right', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            borders.append(border)

        return borders

    def _build_grid_borders(self):
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            borders.append(border)
        return borders

    def _build_no_borders(self):
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            borders.append(border)
        return borders

    def _format_table_autofit(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        layout = tblPr.find(qn('w:tblLayout'))
        if layout is not None:
            tblPr.remove(layout)
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'autofit')
        tblPr.append(layout)

    def _format_table_cell_margins(self, table, margin_cm):
        margin_emu = Cm(margin_cm)
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        existing_margins = tblPr.find(qn('w:tblCellMar'))
        if existing_margins is not None:
            tblPr.remove(existing_margins)

        cell_mar = OxmlElement('w:tblCellMar')
        for side in ['top', 'left', 'bottom', 'right']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:w'), str(margin_emu))
            el.set(qn('w:type'), 'dxa')
            cell_mar.append(el)
        tblPr.append(cell_mar)

    def _format_header_row(self, table, style):
        if not table.rows:
            return
        header_row = table.rows[0]
        for cell in header_row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.alignment = style["alignment"]
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    self._apply_cell_font(run, style["font_name"], style["font_size"], style["bold"])

    def _format_body_rows(self, table, style):
        if len(table.rows) < 2:
            return
        for row in table.rows[1:]:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.alignment = style["alignment"]
                    para.paragraph_format.line_spacing = style["line_spacing"]
                    para.paragraph_format.space_before = Pt(1)
                    para.paragraph_format.space_after = Pt(1)
                    for run in para.runs:
                        self._apply_cell_font(run, style["font_name"], style["font_size"], style["bold"])

    def _apply_cell_font(self, run, font_name, font_size, bold):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)

    def _count_images_in_table(self, table):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    drawings = para._element.findall('.//' + qn('w:drawing'))
                    self.report["images_preserved"] += len(drawings)

    def _ensure_captions_with_tables(self):
        doc_paras = self.doc.paragraphs
        for i, para in enumerate(doc_paras):
            text = para.text.strip()
            if not text or not self._is_table_caption(text):
                continue
            if i + 1 < len(doc_paras):
                next_elem = doc_paras[i + 1]._element
                tbl_siblings = next_elem.findall(qn('w:tbl')) if next_elem.tag == qn('w:body') else []
                next_is_table = (next_elem.tag == qn('w:tbl')) or bool(tbl_siblings)
                if not next_is_table:
                    parent = para._element.getparent()
                    para_idx = list(parent).index(para._element)
                    if para_idx + 1 < len(list(parent)):
                        next_sibling = list(parent)[para_idx + 1]
                        next_is_table = next_sibling.tag == qn('w:tbl')

            if next_is_table:
                pf = para.paragraph_format
                pf.keep_with_next = True

    def save(self, output_path=None):
        if output_path is None:
            output_path = self.docx_path
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"[INFO] Table-processor: saved to {output_path}")


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python table_processor.py <docx_path> [config_path]")
        sys.exit(1)

    docx_path = sys.argv[1]
    config_path = sys.argv[2] if len(sys.argv) > 2 else None

    processor = TableProcessor(docx_path, config_path)
    result = processor.run()
    processor.save()

    print(json.dumps(result, ensure_ascii=False, indent=2))
